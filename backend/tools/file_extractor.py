"""
file_extractor.py — Parse uploaded files and produce a structured extraction
for the Requirements Agent.

Supported types: Excel (.xlsx/.xls), JSON (.json), Word (.docx)

Each extractor returns:
    {
        "file_type": "excel" | "json" | "docx",
        "summary":   str   — natural-language description injected as original_input,
        "preview":   dict  — data shown in the UI preview card,
    }
"""

from __future__ import annotations
import io
import json
import re


# ── Excel ─────────────────────────────────────────────────────────────────────

def _read_xlsx(file_bytes: bytes, filename: str):
    """Read a modern .xlsx (zip-based) workbook via openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb.active
    sheet_name = ws.title or "Sheet1"
    all_rows = list(ws.iter_rows(values_only=True))
    wb.close()
    return all_rows, sheet_name


def _read_xls(file_bytes: bytes, filename: str):
    """Read a legacy .xls (OLE2/BIFF) workbook via xlrd."""
    import xlrd

    book = xlrd.open_workbook(file_contents=file_bytes)
    sheet = book.sheet_by_index(0)
    sheet_name = sheet.name or "Sheet1"
    all_rows = [tuple(sheet.row_values(r)) for r in range(sheet.nrows)]
    return all_rows, sheet_name


def extract_excel(file_bytes: bytes, filename: str) -> dict:
    # Pick the reader by the actual file signature, not just the extension —
    # .xlsx is a zip (starts with "PK"), legacy .xls is OLE2 (\xD0\xCF\x11\xE0).
    # This handles mislabeled files (e.g. a real .xls named .xlsx or vice versa).
    is_ole2 = file_bytes[:4] == b"\xd0\xcf\x11\xe0"
    is_zip = file_bytes[:2] == b"PK"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if is_ole2 or (ext == "xls" and not is_zip):
        reader, kind = _read_xls, "legacy .xls"
    else:
        reader, kind = _read_xlsx, ".xlsx"

    try:
        all_rows, sheet_name = reader(file_bytes, filename)
    except Exception as e:
        return {"error": f"Could not parse '{filename}' as an Excel file ({kind}). Make sure it is a valid .xlsx or .xls file. ({type(e).__name__})"}

    if not all_rows:
        return {"error": "The uploaded Excel file appears to be empty."}

    raw_headers = all_rows[0]
    headers = [
        str(h).strip() if h is not None else f"Column{i + 1}"
        for i, h in enumerate(raw_headers)
    ]
    data_rows = all_rows[1:]
    row_count = len(data_rows)

    def _infer_type(col_idx: int) -> str:
        vals = [
            r[col_idx] for r in data_rows[:20]
            if col_idx < len(r) and r[col_idx] is not None
        ]
        if not vals:
            return "text"
        try:
            for v in vals:
                float(str(v).replace(",", ""))
            return "number"
        except (ValueError, TypeError):
            pass
        date_pat = re.compile(
            r"\d{4}[-/]\d{2}[-/]\d{2}|\d{2}[-/]\d{2}[-/]\d{4}"
        )
        if any(date_pat.search(str(v)) for v in vals):
            return "date"
        return "text"

    col_types = [_infer_type(i) for i in range(len(headers))]

    sample_rows = []
    for row in data_rows[:5]:
        sample_rows.append([
            str(v) if v is not None else ""
            for v in row[:len(headers)]
        ])

    col_desc = ", ".join(f"{h} ({t})" for h, t in zip(headers, col_types))
    summary = (
        f"The user has uploaded an Excel file named '{filename}' "
        f"(sheet: {sheet_name}) with {row_count:,} rows and {len(headers)} columns: {col_desc}.\n\n"
        f"Sample data (first 5 rows):\n"
        + "\n".join(
            "  " + " | ".join(r) for r in sample_rows
        ) +
        f"\n\nPlease extract requirements from this file structure. "
        f"Infer likely KPIs from numeric columns, dimensions from text/date columns "
        f"used for grouping, and the business domain from the column names and context."
    )

    return {
        "file_type": "excel",
        "summary": summary,
        "preview": {
            "sheet_name": sheet_name,
            "columns": headers,
            "col_types": col_types,
            "rows": sample_rows,
            "row_count": row_count,
        },
    }


# ── JSON ──────────────────────────────────────────────────────────────────────

def extract_json(file_bytes: bytes, filename: str) -> dict:
    try:
        data = json.loads(file_bytes.decode("utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError) as e:
        return {"error": f"Could not parse JSON file: {e}"}

    if isinstance(data, list) and data:
        first = data[0]
        headers = list(first.keys()) if isinstance(first, dict) else ["value"]
        sample_rows = []
        for item in data[:5]:
            if isinstance(item, dict):
                sample_rows.append([str(item.get(h, "")) for h in headers])
            else:
                sample_rows.append([str(item)])
        row_count = len(data)
        structure = "array"
    elif isinstance(data, dict):
        headers = list(data.keys())[:20]
        sample_rows = [[str(data.get(h, "")) for h in headers]]
        row_count = 1
        structure = "object"
    else:
        return {"error": "Unsupported JSON structure — expected an array or object."}

    col_desc = ", ".join(headers[:15])
    summary = (
        f"The user has uploaded a JSON file named '{filename}' "
        f"containing a {structure} with {row_count:,} records and fields: {col_desc}.\n\n"
        f"Sample data (first 5 records):\n"
        + "\n".join("  " + " | ".join(r) for r in sample_rows) +
        f"\n\nPlease extract requirements from this structure — infer KPIs, dimensions, "
        f"and business domain from the field names and values."
    )

    return {
        "file_type": "json",
        "summary": summary,
        "preview": {
            "sheet_name": None,
            "columns": headers,
            "col_types": ["text"] * len(headers),
            "rows": sample_rows,
            "row_count": row_count,
        },
    }


# ── Word ──────────────────────────────────────────────────────────────────────

def extract_docx(file_bytes: bytes, filename: str) -> dict:
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        return {"error": f"Could not parse '{filename}' as a Word document. Make sure it is a valid .docx file. ({type(e).__name__})"}

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))

    summary = (
        f"The user has uploaded a Word document named '{filename}' "
        f"containing the following content:\n\n{full_text}"
    )
    if table_lines:
        summary += "\n\nTables found in the document:\n" + "\n".join(table_lines[:20])

    summary += (
        "\n\nPlease extract requirements from this document — identify the use case name, "
        "business domain, data points / metrics, dimensions, data sources, filters, "
        "and any other relevant requirement details mentioned in the text."
    )

    word_count = len(full_text.split())
    preview_text = full_text[:600] + ("…" if len(full_text) > 600 else "")

    return {
        "file_type": "docx",
        "summary": summary,
        "preview": {
            "sheet_name": None,
            "columns": [],
            "col_types": [],
            "rows": [],
            "row_count": 0,
            "text_preview": preview_text,
            "word_count": word_count,
        },
    }


# ── Router ────────────────────────────────────────────────────────────────────

def extract_file(file_bytes: bytes, filename: str) -> dict:
    """Route to the correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("xlsx", "xls"):
        return extract_excel(file_bytes, filename)
    if ext == "json":
        return extract_json(file_bytes, filename)
    if ext == "docx":
        return extract_docx(file_bytes, filename)
    return {
        "error": (
            f"Unsupported file type '.{ext}'. "
            "Please upload an Excel (.xlsx), JSON (.json), or Word (.docx) file."
        )
    }
